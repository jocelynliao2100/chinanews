# streamlit_app.py

import streamlit as st
from docx import Document
from collections import defaultdict, Counter
import matplotlib.pyplot as plt
import re
from datetime import datetime
import io
import pandas as pd
import os

# 嘗試載入 jieba
import jieba
import jieba.analyse

# 嘗試載入 BeautifulSoup，若不存在就設為 None
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None
# 嘗試載入 tqdm，若不存在就用 identity function
try:
    from tqdm import tqdm
except ImportError:
    tqdm = lambda x: x

# --- Font Setup ---
from matplotlib.font_manager import FontProperties

LOCAL_FONT_PATH = "NotoSansCJKtc-Regular.otf"  # 放在同目錄
font_installed = False
font_prop = None
if os.path.exists(LOCAL_FONT_PATH):
    try:
        font_prop = FontProperties(fname=LOCAL_FONT_PATH)
        plt.rcParams["font.family"] = font_prop.get_name()
        plt.rcParams["axes.unicode_minus"] = False
        font_installed = True
    except Exception as e:
        st.warning(f"載入字體失敗，將使用預設字體：{e}")
        plt.rcParams["axes.unicode_minus"] = False
else:
    st.info(f"找不到 {LOCAL_FONT_PATH}，如需中文請自行部署字體。")
    plt.rcParams["axes.unicode_minus"] = False

st.set_page_config(page_title="國台辦新聞稿分析", layout="wide")
st.title("🇨🇳 國台辦新聞稿分析工具")

# ========== Helper Functions ==========

@st.cache_data(show_spinner=False)
def process_uploaded_files(uploaded_files, col_names):
    """
    讀入多個 .docx，解析段落中的日期 (2020-01 ~ 2025-04)，
    回傳 {col_name: { "YYYY-MM": count }} 結構。
    """
    date_re = re.compile(
        r"\b(202[0-5])[-/年\.](0?[1-9]|1[0-2])[-/月\.](0?[1-9]|[12]\d|3[01])日?\b"
    )
    counts = defaultdict(lambda: defaultdict(int))

    for name, uploaded in zip(col_names, uploaded_files):
        data = uploaded.getvalue()
        doc = Document(io.BytesIO(data))
        for para in doc.paragraphs:
            m = date_re.search(para.text)
            if not m:
                continue
            y, mth, _ = m.groups()
            ym = f"{y}-{int(mth):02d}"
            dt = datetime(int(y), int(mth), 1)
            if datetime(2020, 1, 1) <= dt <= datetime(2025, 4, 30):
                counts[name][ym] += 1
    return counts

def plot_counts(counts, col_names, display_names):
    """
    繪製折線圖，並在 Streamlit 中顯示。
    """
    # 先收集所有月份
    all_months = sorted({ym for d in counts.values() for ym in d})
    # 轉成 datetime list，並格式化 labels
    dates = [datetime.strptime(ym, "%Y-%m") for ym in all_months]

    fig, ax = plt.subplots(figsize=(12, 6))
    for orig, disp in zip(col_names, display_names):
        yvals = [counts[orig].get(ym, 0) for ym in all_months]
        if font_installed:
            ax.plot(dates, yvals, marker="o", label=disp, fontproperties=font_prop)
        else:
            ax.plot(dates, yvals, marker="o", label=disp)

    title = "2020.01–2025.04 國台辦各欄目新聞稿數量變化"
    if font_installed:
        ax.set_title(title, fontproperties=font_prop)
        ax.set_xlabel("年月", fontproperties=font_prop)
        ax.set_ylabel("數量", fontproperties=font_prop)
    else:
        ax.set_title(title)
        ax.set_xlabel("年月")
        ax.set_ylabel("數量")

    ax.legend()
    ax.grid(True)
    plt.xticks(rotation=45)
    st.pyplot(fig)

def parse_list_docx(file_bytes):
    """
    從 Word 解析出 [YYYY-MM-DD] 標記的標題與 URL 列表。
    回傳 list of dict(date, title, url)。
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        items = []
        for line in text.splitlines():
            dm = re.search(r"\[(202\d-\d{1,2}-\d{1,2})\]", line)
            if not dm:
                continue
            date_str = dm.group(1)
            tm = re.search(r"\]\s*(.*?)\s*(http", line)
            um = re.search(r"(https?://\S+)", line)
            title = tm.group(1).strip() if tm else ""
            url = um.group(1).strip() if um else ""
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d")
                if datetime(2020,1,1) <= dt <= datetime(2025,4,30):
                    items.append({"date": dt, "title": title, "url": url})
            except:
                continue
        return items
    except Exception as e:
        st.error(f"解析列表失敗：{e}")
        return []

def fetch_content(url):
    """
    簡單用 requests + BeautifulSoup 抓取 class="TRS_Editor" 的文字。
    """
    if BeautifulSoup is None:
        return ""
    try:
        r = requests.get(url, timeout=5)
        r.encoding = r.apparent_encoding
        soup = BeautifulSoup(r.text, "html.parser")
        editor = soup.find("div", class_="TRS_Editor")
        return editor.get_text("\n").strip() if editor else ""
    except Exception as e:
        st.warning(f"抓取 {url} 失敗：{e}")
        return ""

# ========== 主程式 ==========

# 1️⃣ 新聞稿數量變化
st.header("1. 新聞稿數量變化分析")

uploaded_counts = st.file_uploader(
    "一次上傳五個 .docx 檔案 (台辦動態、交流交往、政務要聞、部門涉台、新聞發佈)",
    type="docx", accept_multiple_files=True, key="uploader_counts"
)

display_names = ["台辦動態", "交流交往", "政務要聞", "部門涉台", "新聞發佈"]

if uploaded_counts:
    if len(uploaded_counts) != 5:
        st.warning(f"請上傳五個檔，目前 {len(uploaded_counts)} 個")
    else:
        filenames = [f.name for f in uploaded_counts]
        counts = process_uploaded_files(uploaded_counts, filenames)
        plot_counts(counts, filenames, display_names)

        # 額外指標：最多月 & 該月關鍵字
        # 找出哪一個 YYYY-MM 總和最大
        total = Counter()
        for d in counts.values():
            total.update(d)
        most_month, most_cnt = total.most_common(1)[0]
        st.metric("🎯 新聞稿最多月份", most_month, f"{most_cnt} 篇")

        # 抓出該月所有段落，做前20關鍵詞
        segments = []
        # 重新讀段落以保留文字
        date_re = re.compile(
            r"\b(202[0-5])[-/年\.](0?[1-9]|1[0-2])[-/月\.](0?[1-9]|[12]\d|3[01])日?\b"
        )
        for up in uploaded_counts:
            doc = Document(io.BytesIO(up.getvalue()))
            for p in doc.paragraphs:
                m = date_re.search(p.text)
                if m:
                    ym = f"{m.group(1)}-{int(m.group(2)):02d}"
                    if ym == most_month:
                        segments.append(p.text)

        words = []
        for seg in segments:
            words += [w for w in jieba.lcut(seg) if len(w) > 1]
        top20 = Counter(words).most_common(20)

        st.subheader(f"📑 {most_month} 前20關鍵詞")
        cols = st.columns(2)
        for i, (w, cnt) in enumerate(top20):
            cols[i % 2].write(f"{i+1}. {w}：{cnt}")

# 2️⃣ 單篇關鍵詞分析
st.markdown("---")
st.header("2. 單篇新聞稿關鍵詞分析")

uploaded_kw = st.file_uploader(
    "上傳單一 .docx 進行關鍵詞分析", type="docx", key="uploader_kw"
)
if uploaded_kw:
    data = uploaded_kw.getvalue()
    doc = Document(io.BytesIO(data))
    txt = "\n".join(p.text for p in doc.paragraphs)
    tags = jieba.analyse.extract_tags(txt, topK=30, withWeight=True)
    st.subheader("前30關鍵詞 (含權重)")
    for w, wt in tags:
        st.write(f"{w}：{wt:.3f}")

# 3️⃣ 列表解析與全文爬取
st.markdown("---")
st.header("3. 新聞列表解析與全文爬取")

uploaded_list = st.file_uploader(
    "上傳含列表的 .docx (格式：[YYYY-MM-DD] 標題 http...)", 
    type="docx", key="uploader_list"
)
if uploaded_list:
    data = uploaded_list.getvalue()
    items = parse_list_docx(data)
    if items:
        st.success(f"共解析到 {len(items)} 則新聞")
        # 顯示標題與連結
        for it in items:
            st.write(f"- {it['date'].date()}  [{it['title']}]({it['url']})")
        # 爬取全文
        contents = []
        with st.spinner("爬取中..."):
            for it in tqdm(items):
                contents.append(fetch_content(it["url"]))
        st.success(f"成功爬取 {sum(bool(c) for c in contents)} 篇")
        full_text = "\n\n".join(contents)
        st.subheader("全文摘要前 500 字")
        st.write(full_text[:500] + "...")
    else:
        st.warning("未能解析任何列表項目，請確認格式。")
